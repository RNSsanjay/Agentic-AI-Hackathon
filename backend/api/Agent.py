from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import docx
import json
import os
import logging
import re
from datetime import datetime
from typing import TypedDict, List, Dict, Any
from django.conf import settings
import requests
from urllib.parse import urlparse
import base64

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize LLM with proper error handling
llm = None
try:
    # Try different import methods for compatibility
    try:
        # Import google.generativeai directly to avoid the Modality issue
        import google.generativeai as genai
        
        class DirectGeminiLLM:
            def __init__(self, api_key):
                genai.configure(api_key=api_key)
                self.model = genai.GenerativeModel('gemini-1.5-flash')
            
            def invoke(self, prompt):
                try:
                    response = self.model.generate_content(prompt)
                    return SimpleResponse(response.text)
                except Exception as e:
                    logger.error(f"Gemini API error: {str(e)}")
                    raise e
        
        class SimpleResponse:
            def __init__(self, content):
                self.content = content
                self.text = content
        
        llm = DirectGeminiLLM(settings.GEMINI_API_KEY)
        logger.info("✅ Gemini API initialized with direct Google AI")
        
    except Exception as e:
        logger.error(f"Direct Gemini initialization failed: {str(e)}")
        
        # Fallback to langchain if available
        try:
            from langchain_google_genai import ChatGoogleGenerativeAI
            os.environ["GOOGLE_API_KEY"] = settings.GEMINI_API_KEY
            llm = ChatGoogleGenerativeAI(
                model="gemini-1.5-flash", 
                temperature=0.3,
                convert_system_message_to_human=True
            )
            logger.info("✅ Gemini API initialized with ChatGoogleGenerativeAI")
        except Exception as e2:
            logger.error(f"ChatGoogleGenerativeAI also failed: {str(e2)}")
            llm = None

except Exception as e:
    logger.error(f"Failed to initialize Gemini API: {str(e)}")
    llm = None

# Try to import langgraph, with fallback if not available
try:
    from langgraph.graph import StateGraph, END
    LANGGRAPH_AVAILABLE = True
    logger.info("✅ LangGraph imported successfully")
except ImportError as e:
    logger.warning(f"LangGraph not available: {str(e)}. Using sequential processing.")
    LANGGRAPH_AVAILABLE = False

# Try to import MongoDB service
try:
    from utils.mongodb_service import mongodb_service
    MONGODB_AVAILABLE = True
    logger.info("✅ MongoDB service imported successfully")
except ImportError as e:
    logger.warning(f"MongoDB service not available: {str(e)}")
    MONGODB_AVAILABLE = False
    mongodb_service = None

# Try to import RAG system with fallback
try:
    from utils.rag_system import internship_rag
    RAG_AVAILABLE = True
    logger.info("✅ RAG system imported successfully")
except ImportError as e:
    logger.warning(f"RAG system not available: {str(e)}")
    RAG_AVAILABLE = False
    internship_rag = None

# GitHub Analyzer for profile insights
class GitHubAnalyzer:
    def __init__(self):
        self.github_token = getattr(settings, 'GITHUB_TOKEN', None)
        self.base_url = "https://api.github.com"
    
    def get_github_insights(self, resume_text, github_link=None):
        """Extract GitHub insights from resume or provided link"""
        try:
            # Try to find GitHub username from resume or provided link
            github_username = None
            
            if github_link:
                # Extract username from provided GitHub link
                github_patterns = [
                    r'github\.com/([A-Za-z0-9\-_.]+)',
                    r'https?://github\.com/([A-Za-z0-9\-_.]+)',
                ]
                for pattern in github_patterns:
                    match = re.search(pattern, github_link, re.IGNORECASE)
                    if match:
                        github_username = match.group(1)
                        break
            
            # If no link provided, try to extract from resume
            if not github_username:
                github_patterns = [
                    r'github\.com/([A-Za-z0-9\-_.]+)',
                    r'https?://github\.com/([A-Za-z0-9\-_.]+)',
                ]
                for pattern in github_patterns:
                    match = re.search(pattern, resume_text, re.IGNORECASE)
                    if match:
                        github_username = match.group(1)
                        break
            
            if not github_username:
                logger.info("No GitHub username found in resume or provided link")
                return None
            
            # Clean username (remove any trailing paths)
            github_username = github_username.split('/')[0].split('?')[0].split('#')[0]
            
            # Get GitHub profile data
            profile_data = self._fetch_github_profile(github_username)
            if not profile_data:
                return None
            
            # Get repositories data
            repos_data = self._fetch_github_repos(github_username)
            
            # Analyze the data
            analysis = self._analyze_github_data(profile_data, repos_data)
            
            return {
                'username': github_username,
                'profile_url': f"https://github.com/{github_username}",
                'profile': profile_data,
                'repositories': repos_data,
                'analysis': analysis,
                'github_score': self._calculate_github_score(profile_data, repos_data, analysis)
            }
            
        except Exception as e:
            logger.error(f"GitHub analysis failed: {str(e)}")
            return None
    
    def _fetch_github_profile(self, username):
        """Fetch GitHub profile data"""
        try:
            headers = {}
            if self.github_token:
                headers['Authorization'] = f'token {self.github_token}'
            
            response = requests.get(f"{self.base_url}/users/{username}", headers=headers, timeout=10)
            if response.status_code == 200:
                return response.json()
            else:
                logger.warning(f"GitHub API returned status {response.status_code} for user {username}")
                return None
        except Exception as e:
            logger.error(f"Failed to fetch GitHub profile: {str(e)}")
            return None
    
    def _fetch_github_repos(self, username):
        """Fetch GitHub repositories data"""
        try:
            headers = {}
            if self.github_token:
                headers['Authorization'] = f'token {self.github_token}'
            
            # Get up to 100 repos, sorted by last updated
            response = requests.get(
                f"{self.base_url}/users/{username}/repos?sort=updated&per_page=100", 
                headers=headers, timeout=10
            )
            if response.status_code == 200:
                return response.json()
            else:
                logger.warning(f"GitHub API returned status {response.status_code} for repos of {username}")
                return []
        except Exception as e:
            logger.error(f"Failed to fetch GitHub repos: {str(e)}")
            return []
    
    def _analyze_github_data(self, profile, repos):
        """Analyze GitHub profile and repositories"""
        analysis = {
            'languages': {},
            'technologies': set(),
            'project_types': [],
            'total_stars': 0,
            'total_forks': 0,
            'recent_activity': [],
            'top_repos': []
        }
        
        if not repos:
            return analysis
        
        # Analyze repositories
        for repo in repos:
            # Count stars and forks
            analysis['total_stars'] += repo.get('stargazers_count', 0)
            analysis['total_forks'] += repo.get('forks_count', 0)
            
            # Track languages
            language = repo.get('language')
            if language:
                analysis['languages'][language] = analysis['languages'].get(language, 0) + 1
            
            # Detect technologies from repo names and descriptions
            repo_text = f"{repo.get('name', '')} {repo.get('description', '')}".lower()
            
            # Common technologies to detect
            tech_keywords = {
                'react': 'React', 'vue': 'Vue.js', 'angular': 'Angular',
                'django': 'Django', 'flask': 'Flask', 'express': 'Express.js',
                'node': 'Node.js', 'typescript': 'TypeScript',
                'docker': 'Docker', 'kubernetes': 'Kubernetes',
                'aws': 'AWS', 'azure': 'Azure', 'gcp': 'Google Cloud',
                'mongodb': 'MongoDB', 'postgresql': 'PostgreSQL', 'mysql': 'MySQL',
                'redis': 'Redis', 'elasticsearch': 'Elasticsearch',
                'tensorflow': 'TensorFlow', 'pytorch': 'PyTorch',
                'scikit': 'Scikit-learn', 'pandas': 'Pandas', 'numpy': 'NumPy'
            }
            
            for keyword, tech in tech_keywords.items():
                if keyword in repo_text:
                    analysis['technologies'].add(tech)
            
            # Determine project types
            if any(keyword in repo_text for keyword in ['web', 'frontend', 'react', 'vue', 'angular']):
                if 'Web Development' not in analysis['project_types']:
                    analysis['project_types'].append('Web Development')
            
            if any(keyword in repo_text for keyword in ['api', 'backend', 'server', 'django', 'flask']):
                if 'Backend Development' not in analysis['project_types']:
                    analysis['project_types'].append('Backend Development')
            
            if any(keyword in repo_text for keyword in ['data', 'analysis', 'ml', 'ai', 'machine']):
                if 'Data Science' not in analysis['project_types']:
                    analysis['project_types'].append('Data Science')
            
            if any(keyword in repo_text for keyword in ['mobile', 'android', 'ios', 'flutter', 'react-native']):
                if 'Mobile Development' not in analysis['project_types']:
                    analysis['project_types'].append('Mobile Development')
        
        # Convert technologies set to list
        analysis['technologies'] = list(analysis['technologies'])
        
        # Get top repositories by stars
        sorted_repos = sorted(repos, key=lambda x: x.get('stargazers_count', 0), reverse=True)
        analysis['top_repos'] = sorted_repos[:5]
        
        return analysis
    
    def _calculate_github_score(self, profile, repos, analysis):
        """Calculate GitHub activity score (0-100)"""
        score = 0
        
        if not profile:
            return 0
        
        # Base score from profile metrics
        public_repos = profile.get('public_repos', 0)
        followers = profile.get('followers', 0)
        following = profile.get('following', 0)
        
        # Repository count (0-30 points)
        score += min(30, public_repos * 2)
        
        # Followers/Following ratio (0-20 points)
        if followers > 0:
            score += min(20, followers)
        
        # Stars received (0-20 points)
        total_stars = analysis.get('total_stars', 0)
        score += min(20, total_stars)
        
        # Language diversity (0-15 points)
        language_count = len(analysis.get('languages', {}))
        score += min(15, language_count * 3)
        
        # Technology usage (0-15 points)
        tech_count = len(analysis.get('technologies', []))
        score += min(15, tech_count * 2)
        
        return min(100, score)

# Initialize GitHub analyzer
github_analyzer = GitHubAnalyzer()

# File Manager
class FileManager:
    def save_resume(self, file, user_id=None):
        try:
            file_id = f"resume_{datetime.utcnow().timestamp()}"
            file_path = os.path.join(settings.MEDIA_ROOT, f"resumes/{file_id}.docx")
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, 'wb+') as destination:
                for chunk in file.chunks():
                    destination.write(chunk)
            file_info = {"size": file.size, "name": file.name}
            return file_path, file_id, file_info
        except Exception as e:
            logger.error(f"File save error: {str(e)}")
            raise

    def delete_file(self, file_path):
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            logger.error(f"File deletion error: {str(e)}")

file_manager = FileManager()

# Enhanced logging for agent communication with micro goals
def log_agent_communication(agent_name, step, data=None, success=True, micro_goal=None):
    """Log agent communication with detailed information and micro goals"""
    status = "✅ SUCCESS" if success else "❌ FAILED"
    logger.info(f"🤖 AGENT [{agent_name}] - {step} - {status}")
    
    if micro_goal:
        logger.info(f"   🎯 MICRO GOAL: {micro_goal}")
    
    if data and isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, list):
                logger.info(f"   📊 {key}: {len(value)} items")
            elif isinstance(value, str) and len(value) > 100:
                logger.info(f"   📝 {key}: {value[:100]}...")
            else:
                logger.info(f"   📋 {key}: {value}")
    logger.info("=" * 80)

# State definition for the new workflow
class AnalysisState(TypedDict):
    resume_text: str
    github_link: str
    preferences: List[str]
    student_profile: Dict[str, Any]
    best_fit_internships: List[Dict]
    portfolio_gaps: List[Dict]
    rag_aligned_requirements: Dict[str, Any]
    readiness_evaluations: List[Dict]
    extraction_info: Dict[str, Any]
    processing_timestamp: str
    error: str
    current_step: str
    step_progress: int
    detailed_extraction: Dict[str, Any]
    agent_communications: List[Dict]

# Sequential workflow processor for when LangGraph is not available
class SequentialWorkflow:
    def __init__(self, agents):
        self.agents = agents
    
    def invoke(self, initial_state):
        state = initial_state.copy()
        
        for agent_name, agent_func in self.agents:
            try:
                logger.info(f"🔄 Running agent: {agent_name}")
                state = agent_func(state)
                if state.get('error'):
                    logger.error(f"❌ Agent {agent_name} failed: {state['error']}")
                    break
            except Exception as e:
                state['error'] = f"Agent {agent_name} failed: {str(e)}"
                logger.error(f"❌ Agent {agent_name} exception: {str(e)}")
                break
        
        return state

# Agent 1: Student Profile Analyzer
def student_profile_analyzer(state: AnalysisState) -> AnalysisState:
    agent_name = "STUDENT_PROFILE_ANALYZER"
    
    if "agent_communications" not in state:
        state["agent_communications"] = []
    
    comm_log = {
        "agent": agent_name,
        "step": "INITIALIZATION",
        "micro_goal": "Initialize profile analysis with resume text processing",
        "status": "started",
        "timestamp": datetime.utcnow().isoformat(),
        "data": {}
    }
    state["agent_communications"].append(comm_log)
    
    log_agent_communication(agent_name, "STARTING ANALYSIS", 
                           micro_goal="Extract and analyze student profile from resume")
    
    state["current_step"] = "Analyzing student profile with AI..."
    state["step_progress"] = 20
    
    try:
        resume_text = state["resume_text"]
        
        # Text preprocessing
        comm_log = {
            "agent": agent_name,
            "step": "TEXT_PREPROCESSING",
            "micro_goal": "Extract basic patterns and document statistics",
            "status": "processing",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {
                "text_length": len(resume_text),
                "word_count": len(resume_text.split()),
                "preferences": state["preferences"]
            }
        }
        state["agent_communications"].append(comm_log)
        
        # Extract detailed information
        detailed_extraction = extract_detailed_info(resume_text)
        state["detailed_extraction"] = detailed_extraction
        
        # GitHub Analysis
        github_insights = None
        try:
            comm_log = {
                "agent": agent_name,
                "step": "GITHUB_ANALYSIS",
                "micro_goal": "Extract and analyze GitHub profile from resume and provided link",
                "status": "processing",
                "timestamp": datetime.utcnow().isoformat(),
                "data": {}
            }
            state["agent_communications"].append(comm_log)
            
            # Use provided GitHub link or extract from resume
            github_link = state.get("github_link", "")
            github_insights = github_analyzer.get_github_insights(resume_text, github_link)
            if github_insights:
                logger.info(f"✅ GitHub analysis successful for {github_insights['username']}")
                comm_log = {
                    "agent": agent_name,
                    "step": "GITHUB_ANALYSIS_SUCCESS",
                    "micro_goal": "Successfully analyzed GitHub profile",
                    "status": "success",
                    "timestamp": datetime.utcnow().isoformat(),
                    "data": {
                        "github_username": github_insights['username'],
                        "public_repos": github_insights['profile']['public_repos'],
                        "github_score": github_insights['github_score'],
                        "languages_found": len(github_insights['analysis']['languages']),
                        "technologies_found": len(github_insights['analysis']['technologies'])
                    }
                }
                state["agent_communications"].append(comm_log)
            else:
                logger.info("ℹ️ No GitHub profile found in resume")
                comm_log = {
                    "agent": agent_name,
                    "step": "GITHUB_ANALYSIS_NO_PROFILE",
                    "micro_goal": "No GitHub profile detected in resume",
                    "status": "completed",
                    "timestamp": datetime.utcnow().isoformat(),
                    "data": {"github_found": False}
                }
                state["agent_communications"].append(comm_log)
        except Exception as e:
            logger.warning(f"GitHub analysis failed: {str(e)}")
            comm_log = {
                "agent": agent_name,
                "step": "GITHUB_ANALYSIS_FAILED",
                "micro_goal": "GitHub analysis encountered an error",
                "status": "failed",
                "timestamp": datetime.utcnow().isoformat(),
                "data": {"error": str(e)}
            }
            state["agent_communications"].append(comm_log)
        
        # AI Processing
        if llm:
            try:
                profile = process_with_ai(resume_text, state["preferences"], agent_name, state)
            except Exception as e:
                logger.warning(f"AI processing failed: {str(e)}, using fallback")
                profile = create_enhanced_fallback_profile(resume_text, state["preferences"])
        else:
            logger.warning("LLM not available, using fallback profile creation")
            profile = create_enhanced_fallback_profile(resume_text, state["preferences"])
        
        # Enhance profile with GitHub data
        if github_insights:
            profile = enhance_profile_with_github(profile, github_insights)
        
        # Enhance and validate profile
        profile = enhance_extracted_profile(profile, resume_text, state["preferences"])
        profile = validate_and_enhance_profile(profile, state["preferences"])
        
        state["student_profile"] = profile
        state["extraction_info"] = {
            "chars_extracted": len(resume_text),
            "paragraphs_extracted": resume_text.count('\n\n') + 1,
            "email_found": bool(profile.get("email") and profile["email"] not in ['null', None, '']),
            "phone_found": bool(profile.get("phone") and profile["phone"] not in ['null', None, '']),
            "skills_count": len(profile.get("skills", [])),
            "projects_count": len(profile.get("projects", [])),
            "experience_count": len(profile.get("experience", [])),
            "education_count": len(profile.get("education", [])),
            "sections_found": detailed_extraction["sections_detected"]
        }
        
        # Success log
        comm_log = {
            "agent": agent_name,
            "step": "ANALYSIS_COMPLETED",
            "micro_goal": "Profile analysis completed with full data validation",
            "status": "success",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {
                "profile_completeness": calculate_profile_completeness(profile),
                "extraction_info": state["extraction_info"]
            }
        }
        state["agent_communications"].append(comm_log)
        
        state["current_step"] = "Profile analysis completed successfully"
        state["step_progress"] = 40
        return state
        
    except Exception as e:
        error_msg = str(e)
        comm_log = {
            "agent": agent_name,
            "step": "CRITICAL_FAILURE",
            "micro_goal": "Handle critical failure in profile analysis",
            "status": "failed",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {"error": error_msg}
        }
        state["agent_communications"].append(comm_log)
        
        state["error"] = f"Profile analysis failed: {error_msg}"
        return state

def extract_detailed_info(resume_text):
    """Extract detailed information from resume text"""
    detailed_extraction = {
        "total_characters  ": len(resume_text),
        "total_words": len(resume_text.split()),
        "total_lines": len(resume_text.split('\n')),
        "paragraphs": resume_text.count('\n\n') + 1,
        "email_patterns_found": len(re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)),
        "phone_patterns_found": len(re.findall(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', resume_text)),
        "url_patterns_found": len(re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', resume_text)),
        "sections_detected": []
    }
    
    # Section detection
    section_patterns = {
        'education': [r'\bEDUCATION\b', r'\b(?:education|academic|degree|university|college)\b'],
        'experience': [r'\bEXPERIENCE\b', r'\b(?:experience|employment|work|career|job)\b'],
        'skills': [r'\bSKILLS\b', r'\b(?:skills|technologies|technical|programming)\b'],
        'projects': [r'\bPROJECT[S]?\b', r'\b(?:projects|portfolio|work)\b'],
        'certifications': [r'\bCERTIFICATIONS\b', r'\b(?:certifications?|certificates?)\b'],
        'awards': [r'\bAWARDS\b', r'\b(?:awards?|honors?|achievements?)\b']
    }
    
    for section, patterns in section_patterns.items():
        for pattern in patterns:
            if re.search(pattern, resume_text, re.IGNORECASE):
                detailed_extraction["sections_detected"].append(section.title())
                break
    
    return detailed_extraction

def process_with_ai(resume_text, preferences, agent_name, state):
    """Process resume with AI"""
    prompt_text = f"""
You are an expert resume parser. Extract detailed information from this resume and return ONLY a valid JSON object.

RESUME TEXT:
{resume_text}

Return this EXACT JSON structure:
{{
    "name": "Full person name",
    "email": "email@example.com or null",
    "phone": "phone number or null",
    "location": "location or null",
    "linkedin": "LinkedIn URL or null",
    "github": "GitHub URL (https://github.com/username) or null",
    "website": "website URL or null",
    "summary": "Professional summary text or null",
    "skills": ["List all technical skills found"],
    "programming_languages": ["Python", "Java", etc],
    "frameworks": ["React", "Django", etc],
    "tools": ["Git", "Docker", etc],
    "databases": ["MySQL", "MongoDB", etc],
    "domains": ["Web Development", "Data Science", etc],
    "experience_level": "entry-level OR intermediate OR senior",
    "education": [
        {{
            "degree": "Degree name",
            "institution": "University name",
            "year": "Year or null",
            "gpa": "GPA or null"
        }}
    ],
    "experience": [
        {{
            "title": "Job title",
            "company": "Company name",
            "duration": "Duration",
            "description": "Description"
        }}
    ],
    "projects": [
        {{
            "name": "Project name",
            "description": "Description",
            "technologies": ["Tech used"],
            "duration": "Duration or null",
            "github_url": "GitHub repo URL if mentioned or null"
        }}
    ],
    "certifications": [],
    "achievements": [],
    "languages": ["English"],
    "years_of_experience": "0-1, 1-2, 2-5, 5+"
}}

CRITICAL INSTRUCTIONS:
1. Extract GitHub URLs carefully - look for patterns like "github.com/username" or "https://github.com/username"
2. Include GitHub URLs in projects if mentioned
3. Return ONLY the JSON object, no other text
4. Ensure all URLs are complete and valid
"""
    
    try:
        response = llm.invoke(prompt_text)
        
        # Handle different response formats
        result_text = None
        if hasattr(response, 'content'):
            result_text = response.content
        elif hasattr(response, 'text'):
            result_text = response.text
        else:
            result_text = str(response)
        
        # Clean and parse JSON
        result_text = clean_json_response(result_text)
        profile = json.loads(result_text)
        
        return profile
        
    except Exception as e:
        logger.error(f"AI processing error: {str(e)}")
        raise e

def clean_json_response(text):
    """Clean JSON response from AI"""
    text = text.strip()
    
    # Remove markdown formatting
    if '```json' in text:
        start = text.find('```json') + 7
        end = text.find('```', start)
        if end != -1:
            text = text[start:end].strip()
    elif '```' in text:
        start = text.find('```') + 3
        end = text.rfind('```')
        if end != -1 and end > start:
            text = text[start:end].strip()
    
    # Extract JSON object
    json_match = re.search(r'\{.*\}', text, re.DOTALL)
    if json_match:
        text = json_match.group()
    
    # Fix common JSON issues
    text = re.sub(r',\s*}', '}', text)
    text = re.sub(r',\s*]', ']', text)
    text = text.replace("'", '"')
    text = re.sub(r'\bNone\b', 'null', text)
    text = re.sub(r'\bTrue\b', 'true', text)
    text = re.sub(r'\bFalse\b', 'false', text)
    
    return text

# Agent 2: Internship Matcher - Suggests best-fit internships
def internship_matcher(state: AnalysisState) -> AnalysisState:
    agent_name = "INTERNSHIP_MATCHER"
    
    comm_log = {
        "agent": agent_name,
        "step": "INITIALIZATION",
        "micro_goal": "Find best-fit internships based on student profile",
        "status": "started",
        "timestamp": datetime.utcnow().isoformat(),
        "data": {}
    }
    state["agent_communications"].append(comm_log)
    
    log_agent_communication(agent_name, "STARTING INTERNSHIP MATCHING", 
                           micro_goal="Match student profile with best-fit internships")
    
    state["current_step"] = "Finding best-fit internships..."
    state["step_progress"] = 40
    
    try:
        profile = state["student_profile"]
        preferences = state["preferences"]
        
        # Try RAG system first for intelligent matching
        best_fit_internships = []
        if RAG_AVAILABLE and internship_rag:
            try:
                rag_matches = internship_rag.find_matching_internships(
                    profile=profile,
                    preferences=preferences,
                    top_k=8
                )
                best_fit_internships = process_rag_matches(rag_matches, profile)
                
                comm_log = {
                    "agent": agent_name,
                    "step": "RAG_MATCHING_SUCCESS",
                    "micro_goal": "Successfully matched internships using RAG system",
                    "status": "success",
                    "timestamp": datetime.utcnow().isoformat(),
                    "data": {"rag_matches_found": len(best_fit_internships)}
                }
                state["agent_communications"].append(comm_log)
                
            except Exception as e:
                logger.warning(f"RAG system failed: {str(e)}")
                best_fit_internships = []
                
                comm_log = {
                    "agent": agent_name,
                    "step": "RAG_MATCHING_FAILED",
                    "micro_goal": "RAG matching failed, using fallback",
                    "status": "warning",
                    "timestamp": datetime.utcnow().isoformat(),
                    "data": {"error": str(e)}
                }
                state["agent_communications"].append(comm_log)
        
        # Generate fallback recommendations if needed
        if len(best_fit_internships) < 3:
            fallback_internships = generate_intelligent_recommendations(profile, preferences, best_fit_internships)
            best_fit_internships.extend(fallback_internships)
            
            comm_log = {
                "agent": agent_name,
                "step": "FALLBACK_MATCHING",
                "micro_goal": "Added fallback internships to ensure sufficient matches",
                "status": "success",
                "timestamp": datetime.utcnow().isoformat(),
                "data": {"fallback_matches_added": len(fallback_internships)}
            }
            state["agent_communications"].append(comm_log)
        
        # Limit to best 6 matches and ensure quality
        best_fit_internships = best_fit_internships[:6]
        
        # Enhance internships with additional matching data
        enhanced_internships = []
        for internship in best_fit_internships:
            enhanced = enhance_internship_match(internship, profile)
            enhanced_internships.append(enhanced)
        
        state["best_fit_internships"] = enhanced_internships
        
        comm_log = {
            "agent": agent_name,
            "step": "MATCHING_COMPLETED",
            "micro_goal": "Successfully found and enhanced best-fit internships",
            "status": "success",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {
                "total_matches": len(enhanced_internships),
                "avg_matching_score": sum(i.get('matching_score', 0) for i in enhanced_internships) / len(enhanced_internships) if enhanced_internships else 0
            }
        }
        state["agent_communications"].append(comm_log)
        
        state["current_step"] = "Best-fit internships identified"
        state["step_progress"] = 50
        
        return state
        
    except Exception as e:
        error_msg = str(e)
        comm_log = {
            "agent": agent_name,
            "step": "MATCHING_FAILURE",
            "micro_goal": "Handle failure in internship matching",
            "status": "failed",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {"error": error_msg}
        }
        state["agent_communications"].append(comm_log)
        
        state["error"] = f"Internship matching failed: {error_msg}"
        return state

def enhance_internship_match(internship, profile):
    """Enhance internship with additional matching data"""
    enhanced = internship.copy()
    
    # Calculate detailed skill match
    student_skills = set([s.lower() for s in profile.get('skills', [])])
    required_skills = set([s.lower() for s in internship.get('requirements', [])])
    preferred_skills = set([s.lower() for s in internship.get('preferred_skills', [])])
    
    skill_matches = student_skills.intersection(required_skills)
    preferred_matches = student_skills.intersection(preferred_skills)
    
    # Calculate match percentages
    required_match_percentage = (len(skill_matches) / len(required_skills)) * 100 if required_skills else 100
    preferred_match_percentage = (len(preferred_matches) / len(preferred_skills)) * 100 if preferred_skills else 0
    
    # Domain alignment
    student_domains = [d.lower() for d in profile.get('domains', [])]
    internship_domain = internship.get('domain', '').lower()
    domain_match = any(domain in internship_domain for domain in student_domains)
    
    # Experience level match
    experience_match = profile.get('experience_level') == internship.get('experience_level')
    
    # Add enhancement data
    enhanced.update({
        'skill_matches': list(skill_matches),
        'preferred_matches': list(preferred_matches),
        'required_match_percentage': round(required_match_percentage, 1),
        'preferred_match_percentage': round(preferred_match_percentage, 1),
        'domain_match': domain_match,
        'experience_match': experience_match,
        'overall_fit_score': calculate_overall_fit_score(
            required_match_percentage, preferred_match_percentage, 
            domain_match, experience_match
        ),
        'missing_required_skills': list(required_skills - student_skills),
        'missing_preferred_skills': list(preferred_skills - student_skills)
    })
    
    return enhanced

def calculate_overall_fit_score(required_match, preferred_match, domain_match, experience_match):
    """Calculate overall fit score for an internship"""
    score = 0
    
    # Required skills (40% weight)
    score += (required_match / 100) * 0.4
    
    # Preferred skills (20% weight)
    score += (preferred_match / 100) * 0.2
    
    # Domain match (25% weight)
    score += (1 if domain_match else 0) * 0.25
    
    # Experience match (15% weight)
    score += (1 if experience_match else 0) * 0.15
    
    return round(score * 100, 1)  # Return as percentage

def process_rag_matches(rag_matches, profile):
    """Process RAG matches into recommendation format"""
    recommendations = []
    for i, match in enumerate(rag_matches):
        rec = {
            "id": match.get("id", f"RAG_{i+1}"),
            "title": match.get("title", "Software Engineering Intern"),
            "company": match.get("company", "TechCorp"),
            "domain": match.get("domain", "Technology"),
            "location": match.get("location", "Remote"),
            "duration": match.get("duration", "12 weeks"),
            "stipend": match.get("stipend", "Competitive"),
            "matching_score": match.get("matching_score", 0.75),
            "justification": match.get("justification", "Good skill match"),
            "requirements": match.get("requirements", ["Python", "Communication"]),
            "preferred_skills": match.get("preferred_skills", []),
            "description": match.get("description", ""),
            "experience_level": match.get("experience_level", "entry-level"),
            "application_deadline": match.get("application_deadline", ""),
            "skill_match_percentage": calculate_skill_match(
                profile.get("skills", []), 
                match.get("requirements", [])
            ),
            "domain_match": True
        }
        recommendations.append(rec)
    return recommendations

# Agent 3: Portfolio Gap Detector - Flags missing skills or outdated projects
def portfolio_gap_detector(state: AnalysisState) -> AnalysisState:
    agent_name = "PORTFOLIO_GAP_DETECTOR"
    
    comm_log = {
        "agent": agent_name,
        "step": "INITIALIZATION",
        "micro_goal": "Detect portfolio gaps and missing skills",
        "status": "started",
        "timestamp": datetime.utcnow().isoformat(),
        "data": {}
    }
    state["agent_communications"].append(comm_log)
    
    log_agent_communication(agent_name, "STARTING PORTFOLIO GAP ANALYSIS", 
                           micro_goal="Identify missing skills and outdated projects")
    
    state["current_step"] = "Analyzing portfolio gaps..."
    state["step_progress"] = 60
    
    try:
        profile = state["student_profile"]
        best_fit_internships = state["best_fit_internships"]
        
        # Analyze gaps based on best-fit internships
        gaps = generate_intelligent_gaps(profile, best_fit_internships)
        
        # Add project portfolio analysis
        project_gaps = analyze_project_portfolio(profile, best_fit_internships)
        gaps.extend(project_gaps)
        
        # Add GitHub analysis gaps if available
        github_gaps = analyze_github_gaps(profile)
        gaps.extend(github_gaps)
        
        state["portfolio_gaps"] = gaps
        
        comm_log = {
            "agent": agent_name,
            "step": "GAP_ANALYSIS_COMPLETED",
            "micro_goal": "Portfolio gaps successfully identified",
            "status": "success",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {
                "total_gaps_found": len(gaps),
                "critical_gaps": len([g for g in gaps if g.get('priority') == 'high']),
                "skill_gaps": len([g for g in gaps if g.get('category') == 'technical_skills'])
            }
        }
        state["agent_communications"].append(comm_log)
        
        state["current_step"] = "Portfolio gap analysis completed"
        
        return state
        
    except Exception as e:
        error_msg = str(e)
        comm_log = {
            "agent": agent_name,
            "step": "GAP_ANALYSIS_FAILURE",
            "micro_goal": "Handle failure in portfolio gap detection",
            "status": "failed",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {"error": error_msg}
        }
        state["agent_communications"].append(comm_log)
        
        state["error"] = f"Portfolio gap detection failed: {error_msg}"
        return state

def analyze_project_portfolio(profile, internships):
    """Analyze project portfolio for gaps"""
    project_gaps = []
    projects = profile.get('projects', [])
    project_count = len(projects)
    
    # Check project quantity
    if project_count < 3:
        project_gaps.append({
            "category": "projects",
            "type": "portfolio",
            "title": f"Insufficient Projects ({project_count}/3 minimum)",
            "description": f"Build {3 - project_count} more comprehensive projects to strengthen your portfolio",
            "priority": "high",
            "suggested_action": f"Create {3 - project_count} end-to-end projects showcasing different skills",
            "impact": "Essential for demonstrating practical coding abilities",
            "project_ideas": [
                "Full-stack web application with user authentication",
                "REST API with database integration",
                "Data analysis or visualization project",
                "Mobile-responsive web application",
                "Automation script or tool"
            ]
        })
    
    # Check project diversity
    if projects:
        technologies_used = set()
        for project in projects:
            if isinstance(project, dict) and 'technologies' in project:
                if isinstance(project['technologies'], list):
                    technologies_used.update([tech.lower() for tech in project['technologies']])
        
        if len(technologies_used) < 5:
            project_gaps.append({
                "category": "projects",
                "type": "diversity",
                "title": "Limited Technology Diversity in Projects",
                "description": "Your projects use a limited range of technologies",
                "priority": "medium",
                "suggested_action": "Create projects using different technology stacks",
                "impact": "Demonstrates versatility and adaptability",
                "recommended_technologies": ["React", "Node.js", "Python", "SQL", "Docker"]
            })
    
    # Check for deployment gaps
    deployed_projects = 0
    for project in projects:
        if isinstance(project, dict):
            description = project.get('description', '').lower()
            if any(keyword in description for keyword in ['deployed', 'live', 'hosted', 'production']):
                deployed_projects += 1
    
    if deployed_projects == 0 and project_count > 0:
        project_gaps.append({
            "category": "projects",
            "type": "deployment",
            "title": "No Deployed Projects",
            "description": "None of your projects appear to be deployed or publicly accessible",
            "priority": "medium",
            "suggested_action": "Deploy your best projects using free hosting platforms",
            "impact": "Shows ability to deploy and manage applications",
            "deployment_options": ["Netlify", "Vercel", "Heroku", "GitHub Pages"]
        })
    
    return project_gaps

def analyze_github_gaps(profile):
    """Analyze GitHub presence for gaps"""
    github_gaps = []
    
    github_analysis = profile.get('github_analysis', {})
    github_url = profile.get('github')
    
    if not github_url:
        github_gaps.append({
            "category": "professional_presence",
            "type": "github",
            "title": "No GitHub Profile",
            "description": "Missing GitHub profile to showcase your code and contributions",
            "priority": "high",
            "suggested_action": "Create a professional GitHub profile and upload your projects",
            "impact": "Essential for technical roles - showcases code quality and consistency",
            "github_setup_steps": [
                "Create GitHub account with professional username",
                "Add profile picture and bio",
                "Upload your best projects with documentation",
                "Create a comprehensive README for your profile"
            ]
        })
    else:
        github_score = github_analysis.get('github_score', 0)
        public_repos = github_analysis.get('profile', {}).get('public_repos', 0)
        
        if github_score < 50:
            github_gaps.append({
                "category": "professional_presence",
                "type": "github_activity",
                "title": f"Low GitHub Activity Score ({github_score}/100)",
                "description": "Your GitHub profile needs more activity and better organization",
                "priority": "medium",
                "suggested_action": "Increase GitHub activity with more projects and contributions",
                "impact": "Improves technical credibility and demonstrates consistent coding",
                "improvement_areas": [
                    "Add more repositories",
                    "Improve repository documentation",
                    "Increase commit frequency",
                    "Contribute to open-source projects"
                ]
            })
        
        if public_repos < 3:
            github_gaps.append({
                "category": "professional_presence",
                "type": "github_repos",
                "title": f"Few Public Repositories ({public_repos})",
                "description": "Limited number of public repositories to showcase your work",
                "priority": "medium",
                "suggested_action": "Upload more projects to your GitHub profile",
                "impact": "More repositories demonstrate broader experience and skills",
                "repository_suggestions": [
                    "Personal projects showcasing different technologies",
                    "Contributions to open-source projects",
                    "Code samples and algorithms",
                    "Documentation and learning resources"
                ]
            })
    
    return github_gaps

# Agent 4: RAG-Powered Requirement Aligner
def rag_requirement_aligner(state: AnalysisState) -> AnalysisState:
    agent_name = "RAG_REQUIREMENT_ALIGNER"
    
    comm_log = {
        "agent": agent_name,
        "step": "INITIALIZATION",
        "micro_goal": "Align student profile with actual internship requirements using RAG",
        "status": "started",
        "timestamp": datetime.utcnow().isoformat(),
        "data": {}
    }
    state["agent_communications"].append(comm_log)
    
    log_agent_communication(agent_name, "STARTING RAG ALIGNMENT", 
                           micro_goal="Retrieve and align actual internship expectations")
    
    state["current_step"] = "Aligning with real internship requirements..."
    state["step_progress"] = 75
    
    try:
        profile = state["student_profile"]
        best_fit_internships = state["best_fit_internships"]
        preferences = state["preferences"]
        
        # Use RAG system to get detailed requirements alignment
        rag_aligned_data = {
            "alignment_score": 0.0,
            "requirement_matches": [],
            "skill_gaps_detailed": [],
            "industry_expectations": {},
            "recommendation_improvements": []
        }
        
        if RAG_AVAILABLE and internship_rag and best_fit_internships:
            try:
                # Get detailed analysis from RAG system
                rag_matches = internship_rag.find_matching_internships(
                    profile=profile,
                    preferences=preferences,
                    top_k=len(best_fit_internships)
                )
                
                # Analyze alignment for each best-fit internship
                total_alignment = 0
                requirement_matches = []
                skill_gaps_detailed = []
                
                for internship in best_fit_internships:
                    internship_id = internship.get('id')
                    
                    # Find corresponding RAG match
                    rag_match = next((r for r in rag_matches if r.get('id') == internship_id), None)
                    
                    if rag_match:
                        # Calculate detailed alignment
                        student_skills = set([s.lower() for s in profile.get('skills', [])])
                        required_skills = set([s.lower() for s in internship.get('requirements', [])])
                        preferred_skills = set([s.lower() for s in internship.get('preferred_skills', [])])
                        
                        # Skill matching analysis
                        required_match = student_skills.intersection(required_skills)
                        preferred_match = student_skills.intersection(preferred_skills)
                        missing_required = required_skills - student_skills
                        missing_preferred = preferred_skills - student_skills
                        
                        alignment_score = len(required_match) / max(len(required_skills), 1)
                        total_alignment += alignment_score
                        
                        requirement_matches.append({
                            "internship_id": internship_id,
                            "internship_title": internship.get('title'),
                            "company": internship.get('company'),
                            "alignment_score": alignment_score,
                            "required_skills_matched": list(required_match),
                            "preferred_skills_matched": list(preferred_match),
                            "missing_required_skills": list(missing_required),
                            "missing_preferred_skills": list(missing_preferred),
                            "justification": rag_match.get('justification', ''),
                            "rag_score": rag_match.get('matching_score', 0)
                        })
                        
                        # Add to detailed skill gaps
                        for skill in missing_required:
                            skill_gaps_detailed.append({
                                "skill": skill,
                                "importance": "critical",
                                "required_for": [internship.get('title')],
                                "learning_priority": "high"
                            })
                        
                        for skill in missing_preferred:
                            skill_gaps_detailed.append({
                                "skill": skill,
                                "importance": "preferred",
                                "required_for": [internship.get('title')],
                                "learning_priority": "medium"
                            })
                
                # Calculate overall alignment
                avg_alignment = total_alignment / max(len(best_fit_internships), 1)
                
                # Get industry expectations
                industry_expectations = extract_industry_expectations(best_fit_internships, profile)
                
                # Generate improvement recommendations
                recommendation_improvements = generate_alignment_recommendations(
                    requirement_matches, skill_gaps_detailed, profile
                )
                
                rag_aligned_data = {
                    "alignment_score": avg_alignment,
                    "requirement_matches": requirement_matches,
                    "skill_gaps_detailed": consolidate_skill_gaps(skill_gaps_detailed),
                    "industry_expectations": industry_expectations,
                    "recommendation_improvements": recommendation_improvements,
                    "rag_processing_successful": True
                }
                
                comm_log = {
                    "agent": agent_name,
                    "step": "RAG_ALIGNMENT_COMPLETED",
                    "micro_goal": "Successfully aligned requirements using RAG system",
                    "status": "success",
                    "timestamp": datetime.utcnow().isoformat(),
                    "data": {
                        "alignment_score": avg_alignment,
                        "matches_analyzed": len(requirement_matches),
                        "skill_gaps_found": len(rag_aligned_data["skill_gaps_detailed"])
                    }
                }
                state["agent_communications"].append(comm_log)
                
            except Exception as e:
                logger.warning(f"RAG alignment failed: {str(e)}")
                rag_aligned_data["rag_processing_successful"] = False
                rag_aligned_data["error"] = str(e)
        else:
            # Fallback when RAG is not available
            logger.warning("RAG system not available, using fallback alignment")
            rag_aligned_data = generate_fallback_alignment(profile, best_fit_internships, preferences)
        
        state["rag_aligned_requirements"] = rag_aligned_data
        state["current_step"] = "Requirement alignment completed"
        
        return state
        
    except Exception as e:
        error_msg = str(e)
        comm_log = {
            "agent": agent_name,
            "step": "ALIGNMENT_FAILURE",
            "micro_goal": "Handle failure in requirement alignment",
            "status": "failed",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {"error": error_msg}
        }
        state["agent_communications"].append(comm_log)
        
        state["error"] = f"Requirement alignment failed: {error_msg}"
        return state

def extract_industry_expectations(internships, profile):
    """Extract industry expectations from internship listings"""
    expectations = {
        "common_technologies": {},
        "experience_levels": {},
        "soft_skills": {},
        "education_requirements": {},
        "domain_trends": {}
    }
    
    for internship in internships:
        # Technology frequency
        for req in internship.get('requirements', []):
            tech = req.lower()
            expectations["common_technologies"][tech] = expectations["common_technologies"].get(tech, 0) + 1
        
        # Experience level tracking
        exp_level = internship.get('experience_level', 'entry-level')
        expectations["experience_levels"][exp_level] = expectations["experience_levels"].get(exp_level, 0) + 1
        
        # Domain trends
        domain = internship.get('domain', 'General')
        expectations["domain_trends"][domain] = expectations["domain_trends"].get(domain, 0) + 1
    
    return expectations

def generate_alignment_recommendations(requirement_matches, skill_gaps, profile):
    """Generate recommendations to improve alignment"""
    recommendations = []
    
    # Analyze common missing skills
    missing_skills = {}
    for gap in skill_gaps:
        skill = gap["skill"]
        importance = gap["importance"]
        missing_skills[skill] = missing_skills.get(skill, 0) + (2 if importance == "critical" else 1)
    
    # Sort by importance
    sorted_skills = sorted(missing_skills.items(), key=lambda x: x[1], reverse=True)
    
    # Generate top recommendations
    for skill, weight in sorted_skills[:5]:
        priority = "high" if weight >= 3 else "medium" if weight >= 2 else "low"
        recommendations.append({
            "type": "skill_development",
            "skill": skill,
            "priority": priority,
            "weight": weight,
            "action": f"Learn {skill.title()} to improve alignment",
            "impact": f"Will improve matching for {weight} internship(s)",
            "learning_resources": get_skill_learning_resources(skill)
        })
    
    return recommendations

def consolidate_skill_gaps(skill_gaps_detailed):
    """Consolidate similar skill gaps"""
    consolidated = {}
    
    for gap in skill_gaps_detailed:
        skill = gap["skill"]
        if skill in consolidated:
            consolidated[skill]["required_for"].extend(gap["required_for"])
            if gap["importance"] == "critical":
                consolidated[skill]["importance"] = "critical"
        else:
            consolidated[skill] = gap.copy()
    
    return list(consolidated.values())

def generate_fallback_alignment(profile, internships, preferences):
    """Generate fallback alignment when RAG is not available"""
    return {
        "alignment_score": 0.7,  # Default reasonable score
        "requirement_matches": [],
        "skill_gaps_detailed": [],
        "industry_expectations": {
            "common_technologies": {"python": 3, "javascript": 2, "git": 4},
            "experience_levels": {"entry-level": len(internships)},
            "domain_trends": {pref: 1 for pref in preferences}
        },
        "recommendation_improvements": [],
        "rag_processing_successful": False,
        "fallback_used": True
    }

def get_skill_learning_resources(skill):
    """Get learning resources for a specific skill"""
    resources_map = {
        "python": ["Python.org tutorial", "Codecademy Python", "Automate the Boring Stuff"],
        "javascript": ["MDN Web Docs", "freeCodeCamp", "JavaScript.info"],
        "react": ["React.dev", "React Tutorial", "freeCodeCamp React"],
        "git": ["Git Tutorial", "GitHub Learning Lab", "Atlassian Git Tutorials"],
        "sql": ["W3Schools SQL", "SQLite Tutorial", "PostgreSQL Tutorial"],
        "docker": ["Docker Official Tutorial", "Docker for Beginners", "Play with Docker"]
    }
    
    return resources_map.get(skill.lower(), [f"{skill.title()} official documentation", f"{skill.title()} tutorials", "Online courses"])

# Agent 5: Readiness Evaluator - Outputs readiness score and prep plan for each role
def readiness_evaluator(state: AnalysisState) -> AnalysisState:
    agent_name = "READINESS_EVALUATOR"
    
    comm_log = {
        "agent": agent_name,
        "step": "INITIALIZATION",
        "micro_goal": "Calculate readiness scores and generate preparation plans",
        "status": "started",
        "timestamp": datetime.utcnow().isoformat(),
        "data": {}
    }
    state["agent_communications"].append(comm_log)
    
    state["current_step"] = "Evaluating readiness for each role..."
    state["step_progress"] = 90
    
    try:
        profile = state["student_profile"]
        best_fit_internships = state["best_fit_internships"]
        portfolio_gaps = state["portfolio_gaps"]
        rag_aligned_requirements = state["rag_aligned_requirements"]
        
        # Generate readiness evaluations
        readiness_evaluations = []
        
        # Overall evaluation
        overall_score = calculate_readiness_score(profile, portfolio_gaps)
        overall_evaluation = {
            "internship_title": "Overall Internship Readiness",
            "company": "General Assessment", 
            "readiness_score": overall_score,
            "next_steps": generate_next_steps(profile, portfolio_gaps),
            "timeline": estimate_preparation_timeline(portfolio_gaps)
        }
        readiness_evaluations.append(overall_evaluation)
        
        # Individual internship evaluations
        for internship in best_fit_internships[:3]:
            evaluation = {
                "internship_title": internship.get('title', 'Unknown'),
                "company": internship.get('company', 'Unknown'),
                "readiness_score": internship.get('overall_fit_score', overall_score),
                "next_steps": generate_role_specific_steps(internship, profile),
                "timeline": estimate_preparation_timeline(portfolio_gaps)
            }
            readiness_evaluations.append(evaluation)
        
        state["readiness_evaluations"] = readiness_evaluations
        state["processing_timestamp"] = datetime.utcnow().isoformat()
        state["current_step"] = "Analysis completed successfully!"
        state["step_progress"] = 100
        
        comm_log = {
            "agent": agent_name,
            "step": "READINESS_EVALUATION_COMPLETED",
            "micro_goal": "Readiness evaluation completed",
            "status": "success",
            "timestamp": datetime.utcnow().isoformat(),
            "data": {"readiness_score": overall_score}
        }
        state["agent_communications"].append(comm_log)
        
        return state
        
    except Exception as e:
        state["error"] = f"Readiness evaluation failed: {str(e)}"
        return state

def generate_role_specific_steps(internship, profile):
    """Generate role-specific next steps"""
    steps = []
    missing_skills = internship.get('missing_required_skills', [])
    
    for skill in missing_skills[:3]:
        steps.append({
            "category": "Skill Development",
            "priority": "High",
            "action": f"Learn {skill.title()} for {internship.get('title')}",
            "goal": f"Master {skill.title()} to meet role requirements",
            "timeline": "2-3 weeks"
        })
    
    return steps

# Main workflow setup - New Enhanced Flow
if LANGGRAPH_AVAILABLE:
    # Use LangGraph if available
    from langgraph.graph import StateGraph, END
    
    workflow = StateGraph(AnalysisState)
    workflow.add_node("student_profile_analyzer", student_profile_analyzer)
    workflow.add_node("internship_matcher", internship_matcher)
    workflow.add_node("portfolio_gap_detector", portfolio_gap_detector)
    workflow.add_node("rag_requirement_aligner", rag_requirement_aligner)
    workflow.add_node("readiness_evaluator", readiness_evaluator)
    
    # New Enhanced Flow
    workflow.set_entry_point("student_profile_analyzer")
    workflow.add_edge("student_profile_analyzer", "internship_matcher")
    workflow.add_edge("internship_matcher", "portfolio_gap_detector")
    workflow.add_edge("portfolio_gap_detector", "rag_requirement_aligner")
    workflow.add_edge("rag_requirement_aligner", "readiness_evaluator")
    workflow.add_edge("readiness_evaluator", END)
    
    graph = workflow.compile()
else:
    # Use sequential workflow - New Enhanced Flow
    agents = [
        ("student_profile_analyzer", student_profile_analyzer),
        ("internship_matcher", internship_matcher),
        ("portfolio_gap_detector", portfolio_gap_detector),
        ("rag_requirement_aligner", rag_requirement_aligner),
        ("readiness_evaluator", readiness_evaluator)
    ]
    graph = SequentialWorkflow(agents)

class ResumeAnalysisView(APIView):
    def post(self, request):
        try:
            logger.info("🚀 STARTING RESUME ANALYSIS WORKFLOW")
            
            if not llm:
                return Response({
                    "error": "AI service not available. Please check configuration.",
                    "fallback_available": True
                }, status=status.HTTP_503_SERVICE_UNAVAILABLE)
            
            # Get request data
            resume_file = request.FILES.get("resume")
            preferences = json.loads(request.data.get("preferences", "[]"))
            github_link = request.data.get("github_link", "")  # Get GitHub link if provided
            
            if not resume_file:
                return Response({
                    "error": "Resume file is required"
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Save file
            try:
                file_path, file_id, file_info = file_manager.save_resume(resume_file)
            except Exception as e:
                return Response({
                    "error": f"Failed to save resume: {str(e)}"
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Extract text
            try:
                doc = docx.Document(file_path)
                resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                
                # Extract from tables too
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                resume_text += "\n" + cell.text
                
                if not resume_text.strip():
                    file_manager.delete_file(file_path)
                    return Response({
                        "error": "Could not extract text from document"
                    }, status=status.HTTP_400_BAD_REQUEST)
                    
            except Exception as e:
                file_manager.delete_file(file_path)
                return Response({
                    "error": f"Failed to process document: {str(e)}"
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Initialize state
            initial_state = {
                "resume_text": resume_text,
                "github_link": github_link,
                "preferences": preferences,
                "student_profile": {},
                "best_fit_internships": [],
                "portfolio_gaps": [],
                "rag_aligned_requirements": {},
                "readiness_evaluations": [],
                "extraction_info": {},
                "processing_timestamp": "",
                "error": "",
                "current_step": "Starting analysis...",
                "step_progress": 0,
                "detailed_extraction": {},
                "agent_communications": []
            }
            
            # Run workflow
            final_state = graph.invoke(initial_state)
            
            # Clean up
            file_manager.delete_file(file_path)
            
            if final_state.get("error"):
                return Response({
                    "error": final_state["error"]
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Prepare response
            response_data = {
                "student_profile": final_state["student_profile"],
                "internship_recommendations": final_state.get("best_fit_internships", []),
                "portfolio_gaps": final_state["portfolio_gaps"],
                "rag_aligned_requirements": final_state.get("rag_aligned_requirements", {}),
                "readiness_evaluations": final_state["readiness_evaluations"],
                "extraction_info": final_state["extraction_info"],
                "detailed_extraction": final_state["detailed_extraction"],
                "processing_timestamp": final_state["processing_timestamp"],
                "current_step": final_state["current_step"],
                "step_progress": final_state["step_progress"],
                "agent_communications": final_state["agent_communications"],
                "file_info": {
                    "file_id": file_id,
                    "processed": True,
                    "chars_extracted": len(resume_text),
                    "resume_text": resume_text[:3000] + "..." if len(resume_text) > 3000 else resume_text                }
            }
            
            # Schedule MongoDB save for 2 seconds later (async to avoid blocking response)
            import threading
            
            def delayed_save_to_mongodb():
                """Save analysis to MongoDB after a 2-second delay"""
                import time
                time.sleep(2)  # Wait 2 seconds
                
                logger.info("🚀 Starting delayed MongoDB save process...")
                
                if not MONGODB_AVAILABLE or not mongodb_service:
                    logger.warning("⚠️ MongoDB service not available - analysis not saved (delayed save)")
                    return
                
                try:
                    # First, test the connection
                    logger.info("🔍 Testing MongoDB connection...")
                    if not mongodb_service.is_connected():
                        logger.info("🔄 MongoDB not connected, attempting to reconnect...")
                        connection_success = mongodb_service.connect()
                        if not connection_success:
                            logger.error("❌ MongoDB reconnection failed - cannot save analysis")
                            return
                    
                    # Verify connection is working
                    if mongodb_service.is_connected():
                        # Extract user_id from request
                        user_id = request.data.get("user_id", "anonymous")
                        
                        logger.info(f"💾 Saving analysis to MongoDB for user: {user_id}")
                        analysis_id = mongodb_service.save_analysis_result(response_data, user_id)
                        
                        if analysis_id:
                            logger.info(f"✅ Analysis saved to MongoDB with ID: {analysis_id} (delayed save)")
                            
                            # Update dashboard cache with new analysis results
                            readiness_score = None
                            if final_state.get("readiness_evaluations"): 
                                # Calculate average readiness score
                                scores = [eval.get("overall_score", 0) for eval in final_state["readiness_evaluations"]]
                                readiness_score = round(sum(scores) / len(scores)) if scores else 0
                            
                            internship_matches = len(final_state.get("internship_recommendations", []))
                            gaps_detected = len(final_state.get("portfolio_gaps", []))
                            
                            # Import the cache update function
                            try:
                                from .views import update_analysis_cache
                                update_analysis_cache(readiness_score, internship_matches, gaps_detected)
                                logger.info(f"✅ Dashboard cache updated: score={readiness_score}, matches={internship_matches}, gaps={gaps_detected}")
                            except Exception as cache_error:
                                logger.warning(f"⚠️ Cache update failed: {str(cache_error)}")
                            
                        else:
                            logger.warning("⚠️ Failed to save analysis to MongoDB - no ID returned (delayed save)")
                    else:
                        logger.warning("⚠️ MongoDB connection verification failed")
                        
                except Exception as e:
                    logger.error(f"❌ MongoDB delayed save error: {str(e)}")
                    # Try to get connection diagnostics
                    try:
                        diagnostics = mongodb_service.test_connection()
                        logger.error(f"📊 Connection diagnostics: {diagnostics}")
                    except Exception:
                        logger.error("📊 Could not retrieve connection diagnostics")
            
            # Start the delayed save in a background thread
            if MONGODB_AVAILABLE and mongodb_service:
                save_thread = threading.Thread(target=delayed_save_to_mongodb, daemon=True)
                save_thread.start()
                logger.info("� Scheduled MongoDB save for 2 seconds later (background thread started)")
            else:
                logger.warning("⚠️ MongoDB service not available - skipping delayed save")
            
            return Response(response_data, status=status.HTTP_200_OK)
            
        except Exception as e:
            logger.error(f"💥 CRITICAL WORKFLOW FAILURE: {str(e)}")
            return Response({
                "error": str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
def enhance_profile_with_github(profile, github_insights):
    """Enhance profile with GitHub data"""
    if not github_insights:
        return profile
    
    # Update GitHub URL
    profile['github'] = github_insights['profile_url']
    
    # Enhance skills with GitHub languages and technologies
    github_skills = []
    
    # Add programming languages from GitHub
    for lang, count in github_insights['analysis']['languages'].items():
        if lang not in profile.get('programming_languages', []):
            profile.setdefault('programming_languages', []).append(lang)
        if lang not in profile.get('skills', []):
            github_skills.append(lang)
    
    # Add technologies from GitHub
    for tech in github_insights['analysis']['technologies']:
        if tech not in profile.get('skills', []):
            github_skills.append(tech)
    
    # Update skills list
    profile.setdefault('skills', []).extend(github_skills)
    
    # Update domains based on project types
    github_domains = github_insights['analysis']['project_types']
    profile.setdefault('domains', []).extend([d for d in github_domains if d not in profile['domains']])
    
    # Add GitHub-specific information
    profile['github_analysis'] = {
        'username': github_insights['username'],
        'public_repos': github_insights['profile']['public_repos'],
        'followers': github_insights['profile']['followers'],
        'github_score': github_insights['github_score'],
        'top_languages': list(github_insights['analysis']['languages'].keys())[:5],
        'project_types': github_insights['analysis']['project_types'],
        'total_stars': github_insights['analysis']['total_stars'],
        'total_forks': github_insights['analysis']['total_forks']
    }
    
    # Update experience level based on GitHub activity
    if github_insights['github_score'] > 70:
        profile['experience_level'] = 'senior'
    elif github_insights['github_score'] > 40:
        profile['experience_level'] = 'intermediate'
    elif profile.get('experience_level') != 'senior':
        profile['experience_level'] = 'entry-level'
    
    return profile

# Utility functions
def calculate_profile_completeness(profile):
    """Calculate profile completeness percentage"""
    important_fields = ["name", "email", "skills", "education"]
    completed_fields = sum(1 for field in important_fields if profile.get(field))
    return (completed_fields / len(important_fields)) * 100

def calculate_skill_match(student_skills, required_skills):
    """Calculate skill match percentage"""
    if not student_skills or not required_skills:
        return 0
    
    student_skills_lower = [skill.lower() for skill in student_skills]
    required_skills_lower = [skill.lower() for skill in required_skills]
    
    matches = sum(1 for req_skill in required_skills_lower if req_skill in student_skills_lower)
    return (matches / len(required_skills_lower)) * 100

def calculate_readiness_score(profile, gaps):
    """Calculate readiness score"""
    base_score = calculate_profile_completeness(profile) / 100
    gap_penalty = len([g for g in gaps if g.get('priority') == 'high']) * 0.1
    github_bonus = 0
    
    # Add GitHub bonus
    if profile.get('github_analysis'):
        github_score = profile['github_analysis'].get('github_score', 0)
        github_bonus = (github_score / 100) * 0.2  # Up to 20% bonus
    
    return max(0.3, min(0.95, base_score - gap_penalty + github_bonus))

def generate_next_steps(profile, gaps):
    """Generate comprehensive next steps with goals and priorities"""
    next_steps = []
    
    # Categorize gaps by priority and type
    technical_gaps = [g for g in gaps if g.get('category') in ['technical_skills', 'programming_languages', 'frameworks']]
    project_gaps = [g for g in gaps if g.get('category') == 'projects']
    experience_gaps = [g for g in gaps if g.get('category') == 'experience']
    certification_gaps = [g for g in gaps if g.get('category') == 'certifications']
    
    # 1. Technical Skills Enhancement
    if technical_gaps:
        high_priority_tech = [g for g in technical_gaps if g.get('priority') == 'high']
        if high_priority_tech:
            skill_names = [g.get('skill_name', 'Unknown') for g in high_priority_tech[:2]]
            next_steps.append({
                "category": "Technical Skills",
                "priority": "High",
                "action": f"Master {' and '.join(skill_names)}",
                "description": f"Focus on building strong proficiency in {', '.join(skill_names)} through hands-on practice and projects",
                "goal": f"Achieve intermediate-to-advanced level in {', '.join(skill_names)}",
                "timeline": "4-6 weeks",
                "resources": [
                    "Online courses (Coursera, Udemy, freeCodeCamp)",
                    "Practice platforms (LeetCode, HackerRank, Codewars)",
                    "Official documentation and tutorials"
                ],
                "success_metrics": [
                    "Complete 3-5 projects using these technologies",
                    "Contribute to open-source projects",
                    "Build a portfolio project showcasing these skills"
                ]
            })
    
    # 2. Project Portfolio Development
    project_count = len(profile.get('projects', []))
    if project_count < 3 or project_gaps:
        domain_focus = profile.get('domains', ['Web Development'])[0]
        next_steps.append({
            "category": "Portfolio Projects",
            "priority": "High",
            "action": f"Build {3 - project_count} comprehensive {domain_focus.lower()} projects",
            "description": f"Create end-to-end projects that demonstrate your skills in {domain_focus}",
            "goal": "Have 3-5 well-documented projects showcasing different aspects of your expertise",
            "timeline": "6-8 weeks",
            "resources": [
                "Project idea generators and templates",
                "GitHub for version control and showcasing",
                "Deployment platforms (Netlify, Vercel, Heroku)"
            ],
            "success_metrics": [
                "Each project has clean, documented code",
                "Projects are deployed and accessible online",
                "README files explain project purpose and setup",
                "Code demonstrates best practices and clean architecture"
            ]
        })
    
    # 3. GitHub Profile Optimization
    github_score = profile.get('github_analysis', {}).get('github_score', 0)
    if not profile.get('github') or github_score < 50:
        next_steps.append({
            "category": "Professional Presence",
            "priority": "Medium",
            "action": "Optimize GitHub profile and increase activity",
            "description": "Create a compelling GitHub presence that showcases your coding journey and projects",
            "goal": "Achieve a GitHub score of 70+ with active contributions and well-organized repositories",
            "timeline": "2-3 weeks",
            "resources": [
                "GitHub profile README generators",
                "Open-source contribution guides",
                "Git and GitHub documentation"
            ],
            "success_metrics": [
                "Professional GitHub profile with bio and contact info",
                "Regular commits and contributions",
                "Well-organized repositories with proper documentation",
                "Contribution to open-source projects"
            ]
        })
    
    # 4. Certification and Learning
    experience_level = profile.get('experience_level', 'entry-level')
    if experience_level == 'entry-level' and not certification_gaps:
        relevant_certs = []
        user_skills = [skill.lower() for skill in profile.get('skills', [])]
        
        if any(skill in user_skills for skill in ['python', 'data', 'machine learning']):
            relevant_certs.append("Google Data Analytics Certificate")
        if any(skill in user_skills for skill in ['aws', 'cloud', 'azure']):
            relevant_certs.append("AWS Cloud Practitioner")
        if any(skill in user_skills for skill in ['javascript', 'react', 'web']):
            relevant_certs.append("Meta Frontend Developer Certificate")
        
        if relevant_certs:
            next_steps.append({
                "category": "Professional Certification",
                "priority": "Medium",
                "action": f"Pursue {relevant_certs[0]} certification",
                "description": "Gain industry-recognized credentials to validate your skills and knowledge",
                "goal": f"Complete {relevant_certs[0]} and add to your resume",
                "timeline": "8-12 weeks",
                "resources": [
                    "Official certification courses",
                    "Practice exams and study materials",
                    "Study groups and online communities"
                ],
                "success_metrics": [
                    "Pass the certification exam",
                    "Add certificate to LinkedIn and resume",
                    "Apply learned concepts in projects"
                ]
            })
    
    # 5. Networking and Industry Engagement
    if experience_level == 'entry-level':
        next_steps.append({
            "category": "Professional Networking",
            "priority": "Low",
            "action": "Build professional network and industry presence",
            "description": "Connect with industry professionals and stay updated with latest trends",
            "goal": "Establish meaningful connections and increase industry knowledge",
            "timeline": "Ongoing",
            "resources": [
                "LinkedIn for professional networking",
                "Tech meetups and conferences",
                "Industry blogs and newsletters",
                "Discord/Slack tech communities"
            ],
            "success_metrics": [
                "Connect with 20+ industry professionals",
                "Attend 2-3 tech events or webinars monthly",
                "Engage with tech content on LinkedIn",
                "Join relevant tech communities"
            ]
        })
    
    # 6. Interview Preparation
    if len(next_steps) >= 2:  # If there are significant gaps to address
        next_steps.append({
            "category": "Interview Readiness",
            "priority": "Medium",
            "action": "Prepare for technical interviews and internship applications",
            "description": "Practice coding interviews and prepare compelling application materials",
            "goal": "Be confident and prepared for internship interviews",
            "timeline": "4-6 weeks",
            "resources": [
                "LeetCode, HackerRank for coding practice",
                "Mock interview platforms",
                "Resume review services",
                "Behavioral interview preparation guides"
            ],
            "success_metrics": [
                "Solve 50+ coding problems across different topics",
                "Complete 3-5 mock interviews",
                "Prepare answers for common behavioral questions",
                "Optimize resume and cover letter templates"
            ]
        })
    
    # Sort by priority (High -> Medium -> Low)
    priority_order = {'High': 0, 'Medium': 1, 'Low': 2}
    next_steps.sort(key=lambda x: priority_order.get(x['priority'], 3))
    
    return next_steps[:5]  # Return top 5 most important steps

def estimate_preparation_timeline(gaps):
    """Estimate timeline"""
    high_priority_gaps = [g for g in gaps if g.get('priority') == 'high']
    if len(high_priority_gaps) >= 2:
        return "3-4 months preparation"
    elif len(high_priority_gaps) == 1:
        return "2-3 months preparation"
    else:
        return "1-2 months preparation"

def enhance_extracted_profile(profile, resume_text, preferences):
    """Enhance extracted profile"""
    # Extract skills if missing
    if not profile.get('skills'):
        profile['skills'] = extract_comprehensive_skills(resume_text)
    
    # Extract contact info if missing
    if not profile.get('email'):
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-ZaZ0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)
        if email_match:
            profile['email'] = email_match.group()
    
    if not profile.get('phone'):
        phone_match = re.search(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', resume_text)
        if phone_match:
            profile['phone'] = phone_match.group()
    
    # Extract GitHub URL if missing
    if not profile.get('github'):
        github_patterns = [
            r'github\.com/([A-Za-z0-9\-_.]+)',
            r'https?://github\.com/([A-Za-z0-9\-_.]+)',
        ]
        for pattern in github_patterns:
            match = re.search(pattern, resume_text, re.IGNORECASE)
            if match:
                profile['github'] = f"https://github.com/{match.group(1)}"
                break
    
    return profile

def extract_comprehensive_skills(text):
    """Extract skills from text"""
    skill_database = [
        'Python', 'JavaScript', 'Java', 'C++', 'C#', 'HTML', 'CSS', 'React', 
        'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring', 'MySQL', 
 
        'PostgreSQL', 'MongoDB', 'Git', 'Docker', 'AWS', 'Azure', 'TypeScript',
        'PHP', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin', 'Flutter', 'React Native',
        'TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas', 'NumPy', 'Matplotlib',
        'Kubernetes', 'Jenkins', 'CI/CD', 'Redis', 'Elasticsearch', 'GraphQL',
        'REST', 'API', 'Microservices', 'DevOps', 'Linux', 'Unix', 'Bash'
    ]
    
    found_skills = []
    text_lower = text.lower()
    
    for skill in skill_database:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    
    return found_skills

def validate_and_enhance_profile(profile, preferences):
    """Validate and enhance profile"""
    profile.setdefault("name", "Unknown")
    profile.setdefault("skills", ["Python", "JavaScript"])
    profile.setdefault("domains", preferences[:2] if preferences else ["None"])
    profile.setdefault("experience_level", "entry-level")
    
    # Ensure lists
    list_fields = ["skills", "programming_languages", "frameworks", "tools", 
                   "databases", "domains", "education", "experience", "projects"]
    for field in list_fields:
        if not isinstance(profile.get(field), list):
            profile[field] = []
    
    return profile

def create_enhanced_fallback_profile(resume_text, preferences):
    """Create fallback profile"""
    return {
        "name": extract_name_from_text(resume_text),
        "email": extract_email_from_text(resume_text),
        "phone": extract_phone_from_text(resume_text),
        "github": extract_github_url_from_text(resume_text),
        "skills": extract_comprehensive_skills(resume_text),
        "domains": preferences[:2] if preferences else ["None"],
        "experience_level": "entry-level",
        "education": [],
        "experience": [],
        "projects": [],
        "programming_languages": [],
        "frameworks": [],
        "tools": [],
        "databases": []
    }

def extract_name_from_text(text):
    """Extract name from text"""
    lines = text.split('\n')
    first_line = lines[0].strip() if lines else ""
    if first_line and len(first_line.split()) <= 4:
        return first_line
    return "Unknown"

def extract_email_from_text(text):
    """Extract email from text"""
    match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    return match.group() if match else None

def extract_phone_from_text(text):
    """Extract phone from text"""
    match = re.search(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    return match.group() if match else None

def extract_github_url_from_text(text):
    """Extract GitHub URL from text"""
    patterns = [
        r'https?://github\.com/([A-Za-z0-9\-_.]+)',
        r'github\.com/([A-Za-z0-9\-_.]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            username = match.group(1) if pattern.startswith('https?') else match.group(1)
            return f"https://github.com/{username}"
    return None

def generate_intelligent_recommendations(profile, preferences, existing_recs):
    """Generate intelligent recommendations"""
    recommendations = []
    companies = ["TechCorp", "InnovateTech", "FutureSoft", "DataDriven Inc", "CloudSystems", "DevHub"]
    
    for i, domain in enumerate(preferences[:4]):
        if len(recommendations) + len(existing_recs) >= 6:
            break
        
        # Higher scores for profiles with GitHub
        base_score = 0.75
        if profile.get('github_analysis'):
            github_score = profile['github_analysis'].get('github_score', 0)
            base_score += (github_score / 100) * 0.15  # Up to 15% bonus
        
        rec = {
            "id": f"GEN_{i+1}",
            "title": f"{domain} Intern",
            "company": companies[i % len(companies)],
            "domain": domain,
            "location": "Remote" if i % 2 == 0 else "Hybrid",
            "duration": "12 weeks",
            "stipend": "$6000-8000/month",
            "matching_score": max(0.70, min(0.95, base_score - (i * 0.03))),
            "justification": f"Strong match for {domain} preference" + 
                           (" + excellent GitHub portfolio" if profile.get('github_analysis') else ""),
            "requirements": generate_requirements_for_domain(domain),
            "preferred_skills": ["Communication", "Team Work", "Git"],
            "description": f"Hands-on {domain.lower()} internship with real-world projects",
            "experience_level": profile.get("experience_level", "entry-level"),
            "application_deadline": "2024-04-30",
            "skill_match_percentage": calculate_skill_match(
                profile.get("skills", []), 
                generate_requirements_for_domain(domain)
            ),
            "domain_match": True
        }
        recommendations.append(rec)
    
    return recommendations

def generate_requirements_for_domain(domain):
    """Generate requirements for domain"""
    domain_requirements = {
        "Web Development": ["HTML", "CSS", "JavaScript", "React", "Git"],
        "Data Science": ["Python", "SQL", "Pandas", "Statistics", "Git"],
        "Mobile Development": ["React Native", "JavaScript", "Mobile UI", "Git"],
        "Backend Development": ["Python", "APIs", "Databases", "Git"],
        "Machine Learning": ["Python", "Scikit-learn", "TensorFlow", "Git"],
        "Frontend Development": ["HTML", "CSS", "JavaScript", "React", "Git"],
        "Artificial Intelligence": ["Python", "TensorFlow", "PyTorch", "Git"],
        "Cloud Computing": ["AWS", "Docker", "Kubernetes", "Git"],
        "DevOps": ["Docker", "Kubernetes", "CI/CD", "Git"]
    }
    return domain_requirements.get(domain, ["Python", "Communication", "Git"])

def generate_intelligent_gaps(profile, recommendations):
    """Generate intelligent gaps with comprehensive categorization"""
    gaps = []
    
    # Analyze student profile
    student_skills = set([s.lower() for s in profile.get('skills', [])])
    student_languages = set([lang.lower() for lang in profile.get('programming_languages', [])])
    student_frameworks = set([fw.lower() for fw in profile.get('frameworks', [])])
    project_count = len(profile.get('projects', []))
    experience_count = len(profile.get('experience', []))
    
    # Collect all requirements from recommendations
    all_required_skills = set()
    all_required_languages = set()
    all_required_frameworks = set()
    
    for rec in recommendations:
        for req in rec.get('requirements', []):
            req_lower = req.lower()
            if req_lower in ['python', 'java', 'javascript', 'typescript', 'c++', 'go', 'rust', 'php', 'ruby']:
                all_required_languages.add(req_lower)
            elif req_lower in ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'laravel']:
                all_required_frameworks.add(req_lower)
            else:
                all_required_skills.add(req_lower)
    
    # 1. Technical Skills Gaps
    missing_skills = all_required_skills - student_skills
    critical_skills = ['git', 'sql', 'api', 'testing', 'debugging']
    
    for skill in missing_skills:
        priority = "high" if skill in critical_skills else "medium"
        gaps.append({
            "category": "technical_skills",
            "skill_name": skill.title(),
            "type": "skill",
            "title": f"Missing {skill.title()} Skills",
            "description": f"Learn {skill.title()} to match internship requirements and improve technical proficiency",
            "priority": priority,
            "suggested_action": f"Complete hands-on tutorials and build projects using {skill.title()}",
            "impact": f"Increases matching score for {len([r for r in recommendations if skill in [req.lower() for req in r.get('requirements', [])]])} internships",
            "learning_resources": [
                f"Official {skill.title()} documentation",
                f"Interactive {skill.title()} tutorials",
                f"Practice projects incorporating {skill.title()}"
            ]
        })
      # 2. Programming Language Gaps
    missing_languages = all_required_languages - student_languages
    for lang in list(missing_languages)[:2]:  # Top 2 missing languages
        gaps.append({
            "category": "programming_languages",
            "skill_name": lang.title(),
            "type": "programming_language",
            "title": f"Missing {lang.title()} Programming Language",
            "description": f"Learn {lang.title()} programming language to match internship requirements",
            "priority": "medium",
            "suggested_action": f"Complete {lang.title()} tutorials and build practice projects",
            "impact": f"Increases matching score for {len([r for r in recommendations if lang in [req.lower() for req in r.get('requirements', [])]])} internships",
            "learning_resources": [
                f"Official {lang.title()} documentation",
                f"{lang.title()} coding challenges",
                f"Build projects using {lang.title()}"
            ]
        })
    
    # 6. Professional Experience Gaps
    if experience_count == 0:
        gaps.append({
            "category": "experience",
            "type": "experience",
            "title": "Limited Professional Experience",
            "description": "Gain practical experience through internships, freelance work, or volunteer projects",
            "priority": "medium",
            "suggested_action": "Seek part-time positions, freelance projects, or contribute to open-source initiatives",
            "impact": "Provides real-world experience and professional references",
            "experience_options": [
                "Part-time or contract development work",
                "Volunteer for non-profit organizations",
                "Open-source project contributions",
                "Personal client projects",
                "Hackathon participation"
            ]
        })
    
    # 7. Certification Gaps (for entry-level candidates)
    if profile.get('experience_level') == 'entry-level' and not profile.get('certifications'):
        relevant_domains = profile.get('domains', [])
        if relevant_domains:
            primary_domain = relevant_domains[0].lower()
            cert_suggestions = {
                'web development': 'Meta Frontend Developer Professional Certificate',
                'data science': 'Google Data Analytics Professional Certificate',
                'cloud computing': 'AWS Cloud Practitioner',
                'cybersecurity': 'Google Cybersecurity Professional Certificate'
            }
            
            suggested_cert = None
            for domain, cert in cert_suggestions.items():
                if domain in primary_domain:
                    suggested_cert = cert
                    break
            
            if suggested_cert:
                gaps.append({
                    "category": "certifications",
                    "type": "certification",
                    "title": "Missing Industry Certifications",
                    "description": f"Earn {suggested_cert} to validate your skills and knowledge",
                    "priority": "low",
                    "suggested_action": f"Enroll in and complete {suggested_cert}",
                    "impact": "Provides industry-recognized validation of your skills",
                    "certification_benefits": [
                        "Industry-recognized credential",
                        "Structured learning path",
                        "Enhanced resume credibility",
                        "Network access to certified professionals"
                    ]
                })
    
    # Sort gaps by priority and limit to top 6
    priority_order = {'high': 0, 'medium': 1, 'low': 2}
    gaps.sort(key=lambda x: priority_order.get(x['priority'], 3))
    
    return gaps[:6]  # Return top 6 most critical gaps